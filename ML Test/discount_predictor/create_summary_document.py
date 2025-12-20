from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def create_ml_project_summary():
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('Customer Discount Prediction ML Model', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('Complete Development Summary - Phase 1 to Phase 8', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    exec_summary = doc.add_paragraph()
    exec_summary.add_run(
        "This document presents a comprehensive machine learning solution for predicting customer discount percentages. "
        "The project follows an 8-phase systematic approach, from data understanding to advanced optimization techniques. "
        "The final model achieved a cross-validation RMSE of 7.17, successfully predicting discount percentages for 13,319 customers "
        "with realistic predictions ranging from 20% to 53%. The solution employs Random Forest as the primary algorithm, "
        "enhanced with feature engineering, ensemble methods, and advanced optimization techniques."
    )
    
    doc.add_page_break()
    
    # Phase 1: Data Understanding
    doc.add_heading('Phase 1: Data Understanding', level=1)
    
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(
        "Understand the dataset structure, identify target variables, and categorize feature types for effective preprocessing."
    )
    
    doc.add_heading('Key Findings', level=2)
    findings = doc.add_paragraph()
    findings.add_run("• Dataset Shapes: ").bold = True
    findings.add_run("Train (19,978 × 14), Test (13,319 × 13)\n")
    findings.add_run("• Target Variable: ").bold = True
    findings.add_run("'Discount_percentage' (exists only in training data)\n")
    findings.add_run("• Feature Categories Identified:\n")
    findings.add_run("  - Numerical (9): Orders, bills, ratings, coupons\n")
    findings.add_run("  - Categorical (2): Customer category, premium membership\n")
    findings.add_run("  - Date (1): Last order date\n")
    
    doc.add_heading('Data Quality Issues Discovered', level=2)
    issues = doc.add_paragraph()
    issues.add_run("• Missing values across multiple features\n")
    issues.add_run("• Inconsistent categorical encoding ('0' mixed with proper categories)\n")
    issues.add_run("• Date format requiring conversion\n")
    
    # Phase 2: Data Cleaning & Preprocessing
    doc.add_heading('Phase 2: Data Cleaning & Preprocessing', level=1)
    
    doc.add_heading('Missing Value Strategy', level=2)
    strategy = doc.add_paragraph()
    strategy.add_run("Rationale: ").bold = True
    strategy.add_run("Used median for numerical features (robust to outliers) and mode for categorical features.\n")
    strategy.add_run("• Numerical features: Filled with median values\n")
    strategy.add_run("• Categorical features: Filled with mode or 'Unknown'\n")
    strategy.add_run("• Date features: Converted to datetime, then filled engineered features\n")
    
    doc.add_heading('Feature Engineering (Critical for Performance)', level=2)
    engineering = doc.add_paragraph()
    engineering.add_run("Why Feature Engineering Matters: ").bold = True
    engineering.add_run("Raw features often don't capture business logic. Engineered features significantly improved model performance.\n\n")
    
    engineering.add_run("Date Features (Recency Analysis):\n").bold = True
    engineering.add_run("• days_since_last_order: Customer recency (key predictor)\n")
    engineering.add_run("• last_order_month: Seasonal patterns\n")
    engineering.add_run("• last_order_weekday: Day-of-week effects\n\n")
    
    engineering.add_run("Bill Behavior Features:\n").bold = True
    engineering.add_run("• bill_range: Spending variability indicator\n")
    engineering.add_run("• avg_bill: Average spending level\n\n")
    
    engineering.add_run("Business Logic Features:\n").bold = True
    engineering.add_run("• coupon_usage_ratio: Efficiency of coupon campaigns\n")
    engineering.add_run("• satisfaction_score: Combined customer satisfaction index\n")
    
    doc.add_heading('Encoding Strategy', level=2)
    encoding = doc.add_paragraph()
    encoding.add_run("Decision Rationale:\n").bold = True
    encoding.add_run("• Premium_membership: Binary encoding (natural 0/1 mapping)\n")
    encoding.add_run("• Category_of_customers: One-hot encoding (no ordinal relationship)\n")
    encoding.add_run("• Customer_ID: Removed (identifier, not predictive)\n")
    
    doc.add_page_break()
    
    # Phase 3: Exploratory Data Analysis
    doc.add_heading('Phase 3: Exploratory Data Analysis', level=1)
    
    doc.add_heading('Target Distribution Analysis', level=2)
    target_analysis = doc.add_paragraph()
    target_analysis.add_run("Key Insights:\n").bold = True
    target_analysis.add_run("• Mean: 36.72%, Median: 39.00% (left-skewed distribution)\n")
    target_analysis.add_run("• Skewness: -4.997 (highly skewed - important for model selection)\n")
    target_analysis.add_run("• Outliers: 2.3% (manageable level)\n")
    target_analysis.add_run("• Anomalous values: -99% (data quality flags)\n")
    
    doc.add_heading('Feature-Target Correlation Insights', level=2)
    correlation = doc.add_paragraph()
    correlation.add_run("Top Predictive Features (Business Insights):\n").bold = True
    correlation.add_run("1. Category_Inactive (0.157): Inactive customers receive higher discounts (retention strategy)\n")
    correlation.add_run("2. days_since_last_order (0.143): Recency is key - recent customers get different treatment\n")
    correlation.add_run("3. Premium_membership (0.141): Premium status drives discount eligibility\n")
    correlation.add_run("4. Coupons_offered (0.140): Direct relationship between coupons and discounts\n")
    correlation.add_run("5. No_of_orders_placed (0.116): Order history influences discount strategy\n")
    
    doc.add_heading('Business Logic Validation', level=2)
    validation = doc.add_paragraph()
    validation.add_run("The correlation analysis revealed logical business patterns:\n")
    validation.add_run("• Inactive customers get higher discounts (reactivation strategy)\n")
    validation.add_run("• Premium members receive preferential treatment\n")
    validation.add_run("• Recency matters more than frequency for discount allocation\n")
    
    # Phase 4: Model Building Strategy
    doc.add_heading('Phase 4: Model Building Strategy', level=1)
    
    doc.add_heading('Algorithm Selection Rationale', level=2)
    algorithm = doc.add_paragraph()
    algorithm.add_run("Why Tree-Based Models Over Linear Models:\n").bold = True
    algorithm.add_run("• Non-linear relationships in discount allocation\n")
    algorithm.add_run("• Categorical features with complex interactions\n")
    algorithm.add_run("• Robust to outliers (important given -99% anomalies)\n")
    algorithm.add_run("• Handle missing values naturally\n")
    algorithm.add_run("• Provide feature importance insights\n")
    
    doc.add_heading('Model Comparison Results', level=2)
    results = doc.add_paragraph()
    results.add_run("Baseline Models (Linear):\n").bold = True
    results.add_run("• Linear Regression: 9.194 RMSE\n")
    results.add_run("• Ridge: 9.194 RMSE\n")
    results.add_run("• Lasso: 9.655 RMSE\n\n")
    
    results.add_run("Tree-Based Models:\n").bold = True
    results.add_run("• Random Forest: 7.320 RMSE (22% better than linear)\n")
    results.add_run("• LightGBM: 7.194 RMSE\n")
    results.add_run("• CatBoost: 7.383 RMSE\n")
    results.add_run("• XGBoost: 7.524 RMSE\n")
    
    doc.add_heading('Hyperparameter Tuning Strategy', level=2)
    tuning = doc.add_paragraph()
    tuning.add_run("Smart Tuning Approach:\n").bold = True
    tuning.add_run("• Focused on top 2 performing models\n")
    tuning.add_run("• Used RandomizedSearchCV (efficient vs GridSearch)\n")
    tuning.add_run("• Key parameters: n_estimators, max_depth, learning_rate\n")
    tuning.add_run("• Result: Random Forest improved from 7.320 to 7.174 RMSE\n")
    
    doc.add_page_break()
    
    # Phase 5: Model Evaluation
    doc.add_heading('Phase 5: Model Evaluation', level=1)
    
    doc.add_heading('Model Selection Criteria', level=2)
    criteria = doc.add_paragraph()
    criteria.add_run("Multi-Factor Evaluation:\n").bold = True
    criteria.add_run("• Cross-validation RMSE (primary metric)\n")
    criteria.add_run("• Overfitting assessment (train vs validation)\n")
    criteria.add_run("• Feature importance alignment with business logic\n")
    criteria.add_run("• Model stability across folds\n")
    
    doc.add_heading('Final Model: Tuned Random Forest', level=2)
    final_model = doc.add_paragraph()
    final_model.add_run("Performance Metrics:\n").bold = True
    final_model.add_run("• CV RMSE: 7.173 ± 0.120 (excellent consistency)\n")
    final_model.add_run("• Overfitting Ratio: 1.121 (acceptable - slight overfitting)\n")
    final_model.add_run("• Training R²: 0.608 (explains 60.8% of variance)\n")
    
    doc.add_heading('Feature Importance Validation', level=2)
    importance = doc.add_paragraph()
    importance.add_run("Sanity Check Results: 5/6 expected features in top 5\n").bold = True
    importance.add_run("1. Category_Inactive (0.248): Confirms retention strategy\n")
    importance.add_run("2. Coupons_offered (0.171): Validates coupon-discount relationship\n")
    importance.add_run("3. Premium_membership (0.148): Confirms premium treatment\n")
    importance.add_run("4. Category_Active (0.078): Active customer behavior matters\n")
    importance.add_run("5. No_of_orders_placed (0.075): Order history relevance\n")
    
    # Phase 6: Final Model Training
    doc.add_heading('Phase 6: Final Model Training & Prediction', level=1)
    
    doc.add_heading('Production Model Training', level=2)
    production = doc.add_paragraph()
    production.add_run("Full Dataset Utilization:\n").bold = True
    production.add_run("• Trained on 19,519 clean samples (removed 459 anomalous records)\n")
    production.add_run("• Used all 21 engineered features\n")
    production.add_run("• Applied same preprocessing pipeline to test data\n")
    production.add_run("• Generated 13,319 predictions (100% coverage)\n")
    
    doc.add_heading('Prediction Quality Assurance', level=2)
    quality = doc.add_paragraph()
    quality.add_run("Validation Checks:\n").bold = True
    quality.add_run("• Prediction range: 20.1% to 52.8% (realistic business range)\n")
    quality.add_run("• Mean prediction: 38.1% (aligns with training data)\n")
    quality.add_run("• No negative or extreme values\n")
    quality.add_run("• Standard deviation: 7.8% (appropriate variance)\n")
    
    # Phase 7: Submission Perfection
    doc.add_heading('Phase 7: Submission Perfection', level=1)
    
    doc.add_heading('Quality Assurance Framework', level=2)
    qa = doc.add_paragraph()
    qa.add_run("7-Point Validation System:\n").bold = True
    qa.add_run("✓ No NaN values\n")
    qa.add_run("✓ Correct column names (Customer_ID, Discount_percentage)\n")
    qa.add_run("✓ Exact row count (13,319)\n")
    qa.add_run("✓ No duplicate Customer_IDs\n")
    qa.add_run("✓ Reasonable prediction range (0-100%)\n")
    qa.add_run("✓ Proper data types\n")
    qa.add_run("✓ CSV format compliance\n")
    
    doc.add_page_break()
    
    # Phase 8: Score Boosting
    doc.add_heading('Phase 8: Advanced Optimization Techniques', level=1)
    
    doc.add_heading('Ensemble Strategy', level=2)
    ensemble = doc.add_paragraph()
    ensemble.add_run("Multi-Model Approach:\n").bold = True
    ensemble.add_run("• Combined Random Forest, XGBoost, and LightGBM\n")
    ensemble.add_run("• Weighted by inverse CV RMSE (better models get higher weight)\n")
    ensemble.add_run("• Weights: RF (33.4%), XGBoost (33.3%), LightGBM (33.3%)\n")
    ensemble.add_run("• Reduces prediction variance through averaging\n")
    
    doc.add_heading('Feature Selection Optimization', level=2)
    feature_opt = doc.add_paragraph()
    feature_opt.add_run("Intelligent Feature Pruning:\n").bold = True
    feature_opt.add_run("• Removed 2 weak features (importance < 0.01)\n")
    feature_opt.add_run("• Kept 19 strong predictive features\n")
    feature_opt.add_run("• Minimal performance impact (maintains model simplicity)\n")
    
    doc.add_heading('Log-Target Transformation', level=2)
    log_transform = doc.add_paragraph()
    log_transform.add_run("Addressing Distribution Skewness:\n").bold = True
    log_transform.add_run("• Applied log1p transformation to handle skewed target\n")
    log_transform.add_run("• Achieved RMSE: 0.431 on log scale\n")
    log_transform.add_run("• Better handles extreme values and skewness\n")
    log_transform.add_run("• Provides alternative modeling approach\n")
    
    # Technical Implementation
    doc.add_heading('Technical Implementation Details', level=1)
    
    doc.add_heading('Technology Stack', level=2)
    tech = doc.add_paragraph()
    tech.add_run("Core Libraries:\n").bold = True
    tech.add_run("• pandas: Data manipulation and analysis\n")
    tech.add_run("• scikit-learn: Machine learning algorithms and evaluation\n")
    tech.add_run("• XGBoost: Gradient boosting implementation\n")
    tech.add_run("• LightGBM: Microsoft's gradient boosting framework\n")
    tech.add_run("• CatBoost: Yandex's categorical boosting algorithm\n")
    
    doc.add_heading('Model Parameters (Final Configuration)', level=2)
    params = doc.add_paragraph()
    params.add_run("Random Forest Hyperparameters:\n").bold = True
    params.add_run("• n_estimators: 100 (balance between performance and speed)\n")
    params.add_run("• max_depth: 10 (prevents overfitting while capturing complexity)\n")
    params.add_run("• min_samples_leaf: 4 (regularization parameter)\n")
    params.add_run("• min_samples_split: 2 (default, works well for this dataset)\n")
    params.add_run("• random_state: 42 (reproducibility)\n")
    
    # Business Impact & Insights
    doc.add_heading('Business Impact & Insights', level=1)
    
    doc.add_heading('Key Business Findings', level=2)
    business = doc.add_paragraph()
    business.add_run("Customer Segmentation Insights:\n").bold = True
    business.add_run("• Inactive customers are primary targets for high discounts (retention strategy)\n")
    business.add_run("• Premium members receive preferential discount treatment\n")
    business.add_run("• Recency matters more than frequency in discount allocation\n")
    business.add_run("• Coupon campaigns directly correlate with discount percentages\n")
    
    doc.add_heading('Actionable Recommendations', level=2)
    recommendations = doc.add_paragraph()
    recommendations.add_run("Strategic Recommendations:\n").bold = True
    recommendations.add_run("1. Focus retention efforts on inactive customers with targeted high discounts\n")
    recommendations.add_run("2. Leverage premium membership as a discount differentiator\n")
    recommendations.add_run("3. Implement recency-based discount strategies\n")
    recommendations.add_run("4. Optimize coupon campaigns based on predicted discount preferences\n")
    
    # Model Performance Summary
    doc.add_heading('Final Performance Summary', level=1)
    
    performance_table = doc.add_table(rows=6, cols=2)
    performance_table.style = 'Table Grid'
    
    # Header
    hdr_cells = performance_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    # Data
    metrics = [
        ('Cross-Validation RMSE', '7.173 ± 0.120'),
        ('Training R²', '0.608'),
        ('Test Predictions', '13,319 (100% coverage)'),
        ('Prediction Range', '20.1% - 52.8%'),
        ('Model Type', 'Tuned Random Forest')
    ]
    
    for i, (metric, value) in enumerate(metrics, 1):
        row_cells = performance_table.rows[i].cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion = doc.add_paragraph()
    conclusion.add_run(
        "This comprehensive ML solution successfully addresses the customer discount prediction challenge through "
        "systematic data understanding, intelligent feature engineering, and advanced modeling techniques. "
        "The final Random Forest model achieves excellent performance with realistic predictions, providing "
        "actionable insights for business strategy. The ensemble approach offers additional robustness, "
        "making this solution production-ready for real-world deployment."
    )
    
    # Save document
    doc.save('ML_Discount_Prediction_Summary.docx')
    print("✓ Comprehensive summary saved as: ML_Discount_Prediction_Summary.docx")

if __name__ == "__main__":
    create_ml_project_summary()